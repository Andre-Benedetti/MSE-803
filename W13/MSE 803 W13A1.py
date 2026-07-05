import os
import tkinter as tk
from tkinter import filedialog
from groq import Groq
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

class GroqCVClient:
    def __init__(self, api_key):
        self.client = Groq(api_key=api_key)

    def get_analysis_and_revised_cv(self, cv_text):
        """Requests advice and a full revised CV without extra chatty filler."""
        prompt = f"""
        You are an NZ tech career expert. Analyze the provided CV.
        
        PART 1: Provide 3 actionable recommendations for a Software Engineering role in NZ.
        
        ---SEPARATOR---
        
        PART 2: Provide ONLY the text of the revised, improved CV. 
        Do not include any intro, outro, labels, or notes. 
        Start directly with the contact details and end with the last piece of content.
        
        CV Content:
        {cv_text}
        """
        
        completion = self.client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        return completion.choices[0].message.content

class CVManager:
    def save_recommendations_to_pdf(self, text, output_path):
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph("Career Recommendations", styles['Title']), Spacer(1, 12)]
        
        for line in text.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        doc.build(story)

    def save_revised_cv_to_docx(self, text, output_path):
        doc = Document()
        doc.add_heading('Revised CV', 0)
        # Split text into paragraphs to maintain structure in Word
        for paragraph_text in text.split('\n\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
        doc.save(output_path)

if __name__ == "__main__":
    # Your requested key
    API_KEY = "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"
    client = GroqCVClient(API_KEY)
    manager = CVManager()

    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(title="Select your CV", filetypes=[("Docx files", "*.docx")])

    if file_path:
        try:
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            
            print("Processing... please wait.")
            result = client.get_analysis_and_revised_cv(content)
            
            # Robust parsing logic
            if "---SEPARATOR---" in result:
                parts = result.split("---SEPARATOR---", 1)
                recommendations = parts[0].replace("PART 1:", "").strip()
                revised_cv = parts[1].replace("PART 2:", "").strip()
            else:
                recommendations = "Error parsing recommendations."
                revised_cv = result

            base_dir = os.path.dirname(file_path)
            manager.save_recommendations_to_pdf(recommendations, os.path.join(base_dir, "Recommendations.pdf"))
            manager.save_revised_cv_to_docx(revised_cv, os.path.join(base_dir, "Revised_CV.docx"))
            
            print(f"Success! Files saved in: {base_dir}")
        except Exception as e:
            print(f"An error occurred: {e}")
    else:
        print("No file selected.")